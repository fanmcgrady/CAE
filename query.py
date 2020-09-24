import pymongo
from docx.shared import Inches, Cm

from util import init_word
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

myclient = pymongo.MongoClient("mongodb://localhost:27017/")
mydb = myclient["wenshu"]
zbwenshu = mydb["zbwenshu"]
zbwenshu_relationship = mydb["zbwenshu_relationship"]

table = zbwenshu.find().limit(1)
for item in table:
    doc = item["doc"]
    print(doc)

    text = str(doc)

    index1 = text.find('民 事 判 决 书')

    text1 = doc[0: index1]
    text2 = doc[index1: index1 + 9]

    substr = text[index1 + 9:]
    index2 = substr.find("号")
    text3 = substr[:index2 + 1]

    text4 = substr[index2 + 1:]
    index3 = text4.rfind("。")
    text5 = text4[:index3 + 1]
    print(text4)

    last_text = text4[index3 + 1:]
    last_text = last_text.replace("审", "\n审").\
        replace("代理", "\n代理").\
        replace("二〇", "\n二〇").\
        replace("书记", "\n书记")

    print(last_text)

    word = init_word()
    word.add_paragraph(text1).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word.add_paragraph(text2).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word.add_paragraph(text3).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word.add_paragraph(text5).paragraph_format.first_line_indent = 266700
    word.add_paragraph(last_text).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    word.save("1.docx")