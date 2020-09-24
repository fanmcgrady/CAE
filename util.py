import docx
import xlwt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

_MAPPING = (
u'零', u'一', u'二', u'三', u'四', u'五', u'六', u'七', u'八', u'九', u'十', u'十一', u'十二', u'十三', u'十四', u'十五', u'十六', u'十七',
u'十八', u'十九')
_P0 = (u'', u'十', u'百', u'千',)
_S4 = 10 ** 4

def _to_chinese4(num):
    assert (0 <= num and num < _S4)
    if num < 20:
        return _MAPPING[num]
    else:
        lst = []
        while num >= 10:
            lst.append(num % 10)
            num = num / 10
        lst.append(num)
        c = len(lst)  # 位数
        result = u''

        for idx, val in enumerate(lst):
            val = int(val)
            if val != 0:
                result += _P0[idx] + _MAPPING[val]
                if idx < c - 1 and lst[idx + 1] == 0:
                    result += u'零'
        return result[::-1]

def init_word():
    doc = docx.Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    return doc

# 美化判决书doc内容
def pretty_text(word, doc):
    text = str(doc)

    index1 = text.find('民 事 判 决 书')

    text1 = doc[0: index1]
    text2 = doc[index1: index1 + 9]

    substr = text[index1 + 9:]
    index2 = substr.find("号")
    text3 = substr[:index2 + 1]

    text4 = substr[index2 + 1:]
    index3 = text4.rfind("。")
    text5 = text4[:index3 + 1]

    last_text = text4[index3 + 1:]
    last_text = last_text.replace("审 判", "\n审 判"). \
        replace("代理审判员", "\n代理审判员"). \
        replace("二〇", "\n二〇"). \
        replace("书记", "\n书记"). \
        replace("书 记", "\n书 记"). \
        replace("审核", "\n审核"). \
        replace("撰稿", "\n撰稿"). \
        replace("校对", "\n校对"). \
        replace("印刷", "\n印刷"). \
        replace("人民陪审员", "\n人民陪审员")

    word.add_paragraph(text1).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word.add_paragraph(text2).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word.add_paragraph(text3).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word.add_paragraph(text5).paragraph_format.first_line_indent = 266700
    word.add_paragraph(last_text).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    return word

def init_excel():
    # 创建一个workbook 设置编码
    workbook = xlwt.Workbook(encoding='utf-8')
    # 创建一个worksheet
    worksheet = workbook.add_sheet('My Worksheet')
    # 参数对应 行, 列, 值
    worksheet.write(0, 0, label='序号')
    worksheet.write(0, 1, label='案件id')
    worksheet.write(0, 2, label='案件号')
    worksheet.write(0, 3, label='目标法条')
    worksheet.write(0, 4, label='裁判日期')
    worksheet.write(0, 5, label='案由')
    worksheet.write(0, 6, label='审理程序')
    worksheet.write(0, 7, label='省份')
    worksheet.write(0, 8, label='法院名字')
    worksheet.write(0, 9, label='法院层级')
    worksheet.write(0, 10, label='适用程序')
    worksheet.write(0, 11, label='原告性质')
    worksheet.write(0, 12, label='被告性质')
    worksheet.write(0, 13, label='劳动者性别')
    worksheet.write(0, 14, label='劳动者年龄')
    worksheet.write(0, 15, label='劳动者专业法律人士代理情况')
    worksheet.write(0, 16, label='用人单位专业法律人士代理情况')
    worksheet.write(0, 17, label='除《中华人民共和国劳动合同法》以外的法条适用情况')

    return workbook, worksheet
